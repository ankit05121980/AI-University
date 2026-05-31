"""Render a :class:`Book` to a professional DOCX document using python-docx."""
from __future__ import annotations

import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from ..models import Book

ACCENT = RGBColor(0x43, 0x38, 0xCA)
MUTED = RGBColor(0x47, 0x55, 0x69)


def _add_md_paragraph(doc: Document, line: str) -> None:
    """Add a paragraph, honouring simple **bold** markup and bullets."""
    style = None
    text = line
    if line.lstrip().startswith("- "):
        style = "List Bullet"
        text = line.lstrip()[2:]
    elif re.match(r"^\d+\.\s", line.strip()):
        style = "List Number"
        text = re.sub(r"^\d+\.\s", "", line.strip())
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part:
            p.add_run(part)


def _add_block(doc: Document, body: str) -> None:
    for line in body.split("\n"):
        if line.strip():
            _add_md_paragraph(doc, line)


def render_docx(book: Book, path: str) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Cover
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(book.title)
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(book.subtitle)
    rs.italic = True
    rs.font.size = Pt(14)
    rs.font.color.rgb = MUTED
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"{book.category}  |  Level: {book.level}  |  {book.edition}\n"
        f"Authors: {', '.join(book.authors)}\n"
        f"AI-University Press  |  ISBN {book.isbn}  |  v{book.version}  |  {book.published}"
    )
    doc.add_page_break()

    # Copyright
    doc.add_heading("Copyright", level=1)
    doc.add_paragraph(
        f"Copyright \u00a9 {book.published} AI-University Press. All rights reserved. "
        "Published as part of the AI-University Enterprise AI Knowledge Series."
    )
    doc.add_page_break()

    # TOC
    doc.add_heading("Table of Contents", level=1)
    for ch in book.chapters:
        doc.add_paragraph(f"{ch.number}. {ch.title}  (~{ch.estimated_pages} pp.)")
    doc.add_page_break()

    # Front matter
    fm = book.front_matter
    doc.add_heading("Learning Objectives", level=1)
    for o in fm.learning_objectives:
        doc.add_paragraph(o, style="List Bullet")
    for heading, text in [
        ("Executive Summary", fm.executive_summary),
        ("Industry Context", fm.industry_context),
        ("Business Perspective", fm.business_perspective),
        ("Technical Perspective", fm.technical_perspective),
        ("Architecture Perspective", fm.architecture_perspective),
        ("Governance Perspective", fm.governance_perspective),
        ("Security Perspective", fm.security_perspective),
    ]:
        doc.add_heading(heading, level=1)
        _add_block(doc, text)
    doc.add_page_break()

    # Chapters
    fig_no = 0
    for ch in book.chapters:
        doc.add_heading(f"Chapter {ch.number}. {ch.title}", level=1)
        em = doc.add_paragraph()
        em.add_run(ch.summary).italic = True
        for sec in ch.sections:
            doc.add_heading(sec.heading, level=2)
            _add_block(doc, sec.body)
            if sec.heading.startswith("Architecture"):
                for d in ch.diagrams:
                    fig_no += 1
                    doc.add_heading(f"Figure {fig_no}. {d.title} ({d.fmt})", level=3)
                    code_p = doc.add_paragraph()
                    cr = code_p.add_run(d.source)
                    cr.font.name = "Consolas"
                    cr.font.size = Pt(8)
                    cap = doc.add_paragraph()
                    cap.add_run(d.caption).italic = True
        for cs in ch.code_samples:
            doc.add_heading(f"Listing: {cs.title}", level=3)
            code_p = doc.add_paragraph()
            cr = code_p.add_run(cs.code)
            cr.font.name = "Consolas"
            cr.font.size = Pt(8)
        doc.add_heading("Review Questions", level=2)
        for i, q in enumerate(ch.questions, 1):
            doc.add_paragraph(f"{i}. {q.question}")
            for j, opt in enumerate(q.options):
                doc.add_paragraph(f"{chr(ord('A')+j)}. {opt}", style="List Bullet")
            ans = (f"Answer: {chr(ord('A')+q.answer_index)}. {q.explanation}"
                   if q.answer_index >= 0 else f"Guidance: {q.explanation}")
            ap = doc.add_paragraph()
            ap.add_run(ans).bold = True
        doc.add_page_break()

    # Back matter
    bm = book.back_matter
    doc.add_heading("Hands-On Labs", level=1)
    for lab in bm.labs:
        doc.add_heading(lab["title"], level=2)
        _add_block(doc, lab["body"])
    doc.add_heading("Case Studies", level=1)
    for cs in bm.case_studies:
        doc.add_heading(cs["title"], level=2)
        _add_block(doc, cs["body"])
    doc.add_heading("References", level=1)
    for i, r in enumerate(bm.references, 1):
        doc.add_paragraph(f"{i}. {r}")
    doc.add_heading("Glossary", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Term"
    table.rows[0].cells[1].text = "Definition"
    for g in bm.glossary:
        cells = table.add_row().cells
        cells[0].text = g["term"]
        cells[1].text = g["definition"]
    doc.add_heading("Index", level=1)
    doc.add_paragraph(", ".join(bm.index_terms))

    doc.save(path)
